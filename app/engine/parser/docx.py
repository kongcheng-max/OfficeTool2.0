"""DOCX 解析器 — 基于 python-docx"""

from typing import List

from docx import Document as DocxDocument

from engine.parser.base import BaseParser, Chunk


class DOCXParser(BaseParser):
    name = "docx"
    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def parse(self, file_path: str, original_filename: str) -> List[Chunk]:
        chunks: List[Chunk] = []
        doc = DocxDocument(file_path)

        # 段落提取
        for para in doc.paragraphs:
            if para.text.strip():
                section = para.style.name if para.style else "正文"
                chunks.append(Chunk(
                    content=para.text.strip(),
                    metadata={
                        "section": section,
                        "source": original_filename,
                        "parser_name": self.name,
                        "chunk_index": len(chunks),
                    },
                    chunk_type="text",
                ))

        # 表格提取
        for table in doc.tables:
            markdown = self._table_to_markdown(table)
            if markdown.strip():
                chunks.append(Chunk(
                    content=markdown,
                    metadata={
                        "source": original_filename,
                        "is_table": True,
                        "parser_name": self.name,
                        "chunk_index": len(chunks),
                    },
                    chunk_type="table",
                ))

        return chunks

    @staticmethod
    def _table_to_markdown(table) -> str:
        """将 python-docx 表格转为 Markdown"""
        rows = table.rows
        if not rows:
            return ""

        lines = []
        # 表头（第一行）
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("|" + "|".join("---" for _ in header_cells) + "|")

        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)
